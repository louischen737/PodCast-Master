import os
from typing import Dict, List, Optional
from docx import Document
from ..exceptions import NoValidContentError
from .base_extractor import BaseExtractor

class DocExtractor(BaseExtractor):
    """Word文档内容提取器"""
    
    def extract(self, file_path: str) -> Dict:
        """
        从Word文档中提取内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            Dict: 包含提取内容的字典
        """
        try:
            doc = Document(file_path)
            
            # 提取元数据
            metadata = self._extract_metadata(doc)
            
            # 提取内容
            content = self._extract_content(doc)
            
            if not content:
                raise NoValidContentError("无法从Word文档中提取有效内容")
            
            # 提取关键点
            key_points = self._extract_key_points(content)
            
            return {
                "source_type": "file",
                "title": metadata.get("title", os.path.basename(file_path)),
                "metadata": metadata,
                "content": content,
                "key_points": key_points,
                "cautions": "Word文档可能包含图片、表格和格式化内容，当前仅提取文本内容"
            }
            
        except Exception as e:
            raise NoValidContentError(f"Word文档处理失败: {str(e)}")
    
    def _extract_metadata(self, doc: Document) -> Dict:
        """
        提取Word文档元数据
        
        Args:
            doc: Document实例
            
        Returns:
            Dict: 包含元数据的字典
        """
        metadata = {
            "author": "",
            "created": "",
            "modified": "",
            "paragraph_count": len(doc.paragraphs),
            "section_count": len(doc.sections)
        }
        
        # 尝试从文档属性中获取更多信息
        core_properties = doc.core_properties
        if core_properties:
            metadata.update({
                "author": core_properties.author or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
                "title": core_properties.title or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or ""
            })
        
        return metadata
    
    def _extract_content(self, doc: Document) -> List[Dict]:
        """
        提取Word文档内容
        
        Args:
            doc: Document实例
            
        Returns:
            List[Dict]: 包含内容的列表
        """
        content = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # 检查是否是标题
            if para.style.name.startswith('Heading'):
                current_section = text
                content.append({
                    "section_title": text,
                    "text": "",
                    "content_type": "heading",
                    "level": int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                })
            else:
                # 如果是普通段落
                if current_section is None:
                    current_section = "未命名章节"
                    content.append({
                        "section_title": current_section,
                        "text": text,
                        "content_type": "paragraph"
                    })
                else:
                    # 添加到当前章节
                    content.append({
                        "section_title": current_section,
                        "text": text,
                        "content_type": "paragraph"
                    })
        
        # 处理表格
        for table in doc.tables:
            table_content = []
            for row in table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cell.text.strip())
                table_content.append(row_content)
            
            if table_content:
                content.append({
                    "section_title": current_section or "未命名章节",
                    "text": table_content,
                    "content_type": "table"
                })
        
        return content
    
    def _extract_key_points(self, content: List[Dict]) -> List[str]:
        """
        从内容中提取关键点
        
        Args:
            content: 内容列表
            
        Returns:
            List[str]: 关键点列表
        """
        return [] # 将关键点提取逻辑清空，由AI模块负责 